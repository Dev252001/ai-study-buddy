"""
DOCX text extractor using python-docx.
Extracts body paragraphs, table cells, and heading text.
"""
from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Rough average characters per page for page-count estimation
_CHARS_PER_PAGE_ESTIMATE = 2500


class DOCXProcessor:
    def extract_text(self, file_path: str) -> tuple[str, int]:
        """
        Return *(text, page_estimate)*.
        python-docx does not expose rendered page count, so we estimate it from
        character count.
        """
        from docx import Document
        from docx.oxml.ns import qn

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")

        doc = Document(file_path)
        parts: list[str] = []

        # Paragraphs (includes headings, body text, list items)
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                parts.append(stripped)

        # Table cells
        for table in doc.tables:
            for row in table.rows:
                row_cells: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    parts.append(" | ".join(row_cells))

        text = "\n".join(parts)
        char_count = len(text)
        page_estimate = max(1, round(char_count / _CHARS_PER_PAGE_ESTIMATE))
        return text, page_estimate

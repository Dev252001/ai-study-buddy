"""
Export service — generates PDF, DOCX, and Markdown bytes from content.
Uses reportlab for PDF and python-docx for DOCX.
"""
from __future__ import annotations

import io
import textwrap
from datetime import datetime
from typing import Any

# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------
def _build_pdf(title: str, body_lines: list[tuple[str, str]]) -> bytes:
    """
    Render a clean PDF using ReportLab.

    *body_lines* is a list of (style, text) tuples where style is one of:
      "h1", "h2", "body", "bullet", "code"
    """
    from reportlab.lib import colors  # noqa: PLC0415
    from reportlab.lib.enums import TA_LEFT  # noqa: PLC0415
    from reportlab.lib.pagesizes import A4  # noqa: PLC0415
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # noqa: PLC0415
    from reportlab.lib.units import cm  # noqa: PLC0415
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # noqa: PLC0415

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    base = getSampleStyleSheet()
    styles = {
        "h1": ParagraphStyle("H1", parent=base["Heading1"], fontSize=18, spaceAfter=10),
        "h2": ParagraphStyle("H2", parent=base["Heading2"], fontSize=14, spaceAfter=6),
        "body": ParagraphStyle("Body", parent=base["Normal"], fontSize=11, leading=16),
        "bullet": ParagraphStyle(
            "Bullet", parent=base["Normal"], fontSize=11, leading=16,
            leftIndent=20, bulletIndent=10,
        ),
        "code": ParagraphStyle(
            "Code", parent=base["Code"], fontSize=9, leading=12,
            backColor=colors.HexColor("#f5f5f5"), leftIndent=10, rightIndent=10,
        ),
    }

    story = [
        Paragraph(title, styles["h1"]),
        Paragraph(
            f"Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
            ParagraphStyle("Meta", parent=base["Normal"], fontSize=9, textColor=colors.grey),
        ),
        Spacer(1, 0.5 * cm),
    ]

    for style_name, text in body_lines:
        style = styles.get(style_name, styles["body"])
        # Escape HTML special chars for ReportLab
        safe_text = (
            text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
        )
        if style_name == "bullet":
            safe_text = f"• {safe_text}"
        story.append(Paragraph(safe_text, style))
        story.append(Spacer(1, 0.15 * cm))

    doc.build(story)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------
def _build_docx(title: str, body_lines: list[tuple[str, str]]) -> bytes:
    from docx import Document  # noqa: PLC0415
    from docx.shared import Inches, Pt  # noqa: PLC0415

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(
        f"Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
    ).runs[0].font.size = Pt(9)

    _HEADING_MAP = {"h1": 1, "h2": 2}

    for style_name, text in body_lines:
        if style_name in _HEADING_MAP:
            doc.add_heading(text, level=_HEADING_MAP[style_name])
        elif style_name == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif style_name == "code":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Markdown helper
# ---------------------------------------------------------------------------
def _build_markdown(title: str, body_lines: list[tuple[str, str]]) -> bytes:
    lines = [f"# {title}", f"*Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}*", ""]
    _MD_MAP = {"h1": "## ", "h2": "### ", "bullet": "- ", "code": "    ", "body": ""}
    for style_name, text in body_lines:
        prefix = _MD_MAP.get(style_name, "")
        lines.append(f"{prefix}{text}")
    lines.append("")
    return "\n".join(lines).encode("utf-8")


# ---------------------------------------------------------------------------
# Content → body_lines converters
# ---------------------------------------------------------------------------
def _text_to_lines(content: str) -> list[tuple[str, str]]:
    """Convert plain text / markdown-ish content to (style, text) pairs."""
    lines: list[tuple[str, str]] = []
    for raw_line in content.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            lines.append(("h2", line[3:]))
        elif line.startswith("# "):
            lines.append(("h1", line[2:]))
        elif line.startswith(("- ", "* ", "• ")):
            lines.append(("bullet", line[2:]))
        else:
            lines.append(("body", line))
    return lines


# ---------------------------------------------------------------------------
# Public ExportService
# ---------------------------------------------------------------------------
class ExportService:
    # ------------------------------------------------------------------
    # Generic content exports
    # ------------------------------------------------------------------
    def export_to_pdf(self, content: str, title: str) -> bytes:
        return _build_pdf(title, _text_to_lines(content))

    def export_to_docx(self, content: str, title: str) -> bytes:
        return _build_docx(title, _text_to_lines(content))

    def export_to_markdown(self, content: str, title: str) -> bytes:
        return _build_markdown(title, _text_to_lines(content))

    # ------------------------------------------------------------------
    # Quiz results PDF
    # ------------------------------------------------------------------
    def export_quiz_pdf(self, quiz: Any, attempt: Any) -> bytes:
        lines: list[tuple[str, str]] = [
            ("h2", f"Score: {attempt.percentage:.1f}%  ({int(attempt.score)}/{int(attempt.max_score)})"),
            ("body", f"Type: {quiz.quiz_type}   Difficulty: {quiz.difficulty}"),
            ("body", ""),
            ("h2", "Questions & Answers"),
        ]

        for idx, q in enumerate(quiz.questions, start=1):
            lines.append(("h2", f"Q{idx}: {q.question_text}"))
            if q.options:
                for opt in q.options:
                    lines.append(("bullet", opt))
            q_id = str(q.id)
            user_ans = (attempt.answers or {}).get(q_id, "(no answer)")
            correct = q.correct_answer
            is_correct = user_ans.strip().lower() == correct.strip().lower()
            lines.append(("body", f"Your answer: {user_ans}  {'✓' if is_correct else '✗'}"))
            lines.append(("body", f"Correct answer: {correct}"))
            if q.explanation:
                lines.append(("body", f"Explanation: {q.explanation}"))

        return _build_pdf(quiz.title, lines)

    # ------------------------------------------------------------------
    # Flashcard set PDF
    # ------------------------------------------------------------------
    def export_flashcards_pdf(self, flashcard_set: Any, cards: list[Any]) -> bytes:
        lines: list[tuple[str, str]] = [
            ("body", f"Total cards: {len(cards)}   Difficulty: {flashcard_set.difficulty}"),
            ("body", ""),
        ]
        for idx, card in enumerate(cards, start=1):
            lines.append(("h2", f"Card {idx}"))
            lines.append(("body", f"Front: {card.front}"))
            lines.append(("body", f"Back:  {card.back}"))
            if card.hint:
                lines.append(("body", f"Hint:  {card.hint}"))
            if card.topic:
                lines.append(("body", f"Topic: {card.topic}"))

        return _build_pdf(flashcard_set.title, lines)
